#!/usr/bin/env python3
"""Собирает docs/TZ_e-PTW_Firebase_v1.0.docx из Markdown (упрощённый разбор)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def strip_md_bold(s: str) -> str:
    return s.replace("**", "")


def add_paragraph_with_bold(doc: Document, text: str):
    """Простой разбор **жирного** в одной строке."""
    p = doc.add_paragraph()
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if i % 2 == 1:
            run.bold = True
    return p


def main():
    base = Path(__file__).resolve().parent
    md_path = base / "TZ_e-PTW_Firebase_v1.0.md"
    out_path = base / "TZ_e-PTW_Firebase_v1.0.docx"

    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            t = strip_md_bold(stripped[2:].strip())
            h = doc.add_heading(t, level=0)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif stripped.startswith("## "):
            doc.add_heading(strip_md_bold(stripped[3:].strip()), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(strip_md_bold(stripped[4:].strip()), level=2)
        elif stripped.startswith("|") and "---" not in stripped:
            rows = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                row_line = lines[j].strip()
                if "---" in row_line or row_line.replace("|", "").replace("-", "").replace(" ", "") == "":
                    j += 1
                    continue
                cells = [c.strip() for c in row_line.split("|")[1:-1]]
                if cells:
                    rows.append(cells)
                j += 1
            if rows:
                ncols = max(len(r) for r in rows)
                t = doc.add_table(rows=len(rows), cols=ncols)
                t.style = "Table Grid"
                for ri, r in enumerate(rows):
                    for ci in range(ncols):
                        val = r[ci] if ci < len(r) else ""
                        t.rows[ri].cells[ci].text = strip_md_bold(val)
            i = j
            continue
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_to_paragraph(p, stripped[2:].strip())
        elif stripped.startswith("*") and stripped.endswith("*") and len(stripped) > 2 and not stripped.startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(strip_md_bold(stripped[1:-1].strip()))
            r.italic = True
        else:
            if stripped.startswith("*") and not stripped.startswith("**"):
                p = doc.add_paragraph()
                r = p.add_run(strip_md_bold(stripped[1:].strip() if stripped.endswith("*") else stripped))
                r.italic = True
            else:
                p = doc.add_paragraph()
                add_inline_to_paragraph(p, stripped)
        i += 1

    doc.save(out_path)
    print(f"Written: {out_path}")


def add_inline_to_paragraph(p, text: str):
    if "**" not in text:
        p.add_run(strip_md_bold(text))
        return
    parts = text.split("**")
    for j, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if j % 2 == 1:
            run.bold = True


if __name__ == "__main__":
    main()
